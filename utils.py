import argparse
import os.path
from pathlib import Path
from typing import Dict, Tuple, List, Callable
import requests
import json
import re
from docx import Document
from tqdm import tqdm
from time import time, strftime, gmtime


def time_cost(func: Callable) -> Callable:
    """Descriptor of time cost.

    :param func: Callable function.
    :return: Wrapped function.
    """

    def wrapper(*args, **kwargs):
        """Wrapper of input function.

        :param args: Arguments of input function.
        :param kwargs: Arguments of input function.
        """

        start = time()
        func(*args, **kwargs)
        end = time()
        print(f"Total time: {strftime('%H:%M:%S', gmtime(end - start))}")

    return wrapper


@time_cost
def lazy_export(args: argparse.PARSER) -> None:
    """Lazy export up-to-date paper abstracts of all conferences.

    :param args: Arguments from CLI.
    """

    def _export(_args: argparse.PARSER, _years: List[str]) -> None:
        """Handle different years of each conference.

        :param _args: Adapted Arguments from CLI.
        :param _years: Adapted years.
        """

        for _year in _years:
            _args.year = _year
            parse(_args)

    conferences = ["CVPR", "ICCV", "NeurIPS", "ICLR", "ICML", "ECCV"]
    for conference in conferences:
        args.conference = conference
        if conference == "CVPR":
            _export(args, ["2023", "2022", "2021"])
        elif conference == "ICCV":
            _export(args, ["2023", "2021"])
        elif conference == "NeurIPS":
            _export(args, ["2023", "2022", "2021"])
        elif conference == "ICLR":
            _export(args, ["2024", "2023", "2022", "2021"])
        elif conference == "ICML":
            _export(args, ["2023"])
        elif conference == "ECCV":
            _export(args, ["2022", "2020"])


@time_cost
def parse(args: argparse.PARSER) -> None:
    """Parse requirement.

    :param args: Arguments from CLI.
    """

    source, accepted_types = collate(args.conference, args.year, args.select)
    if source == 'openreview':
        parse_openreview(args.conference, args.year, accepted_types, args.toc, args.pdf, args.wps)
    elif source == 'cvf':
        parse_cvf(args.conference, args.year, args.toc, args.pdf, args.wps)


def collate(conference: str, year: str, select: bool) -> Tuple[str, List[str]]:
    """Collate input conference and corresponding accepted types.

    :param conference: Exact abbreviation of conference.
    :param year: Year of conference.
    :return: Conference source and accepted types.
    :param select: Indicates that whether interactively select specific accepted types.
    """

    openreview = ["NeurIPS", "ICLR", "ICML"]
    cvf = ["CVPR", "ICCV"]
    other = ["ECCV"]
    assert conference in [*openreview, *cvf, *other], "unsupported conference"

    accepted_types = None
    source = None

    if conference in openreview:
        source = "openreview"
        if conference == "NeurIPS":
            assert year in ["2023", "2022", "2021"], f"unsupported year for {conference} so far"

            if year == "2023":
                accepted_types = ["oral", "spotlight", "poster"]
            elif year == "2022":
                accepted_types = ["Accept"]
            elif year == "2021":
                accepted_types = ["Oral", "Spotlight", "Poster"]

        elif conference == "ICLR":
            assert year in ["2024", "2023", "2022", "2021"], f"unsupported year for {conference} so far"

            if year == "2024":
                accepted_types = ["oral", "spotlight", "poster"]
            elif year == "2023":
                accepted_types = ["notable_top_5%", "notable_top_25%", "poster"]
            elif year in ["2022", "2021"]:
                accepted_types = ["Oral", "Spotlight", "Poster"]

        elif conference == "ICML":
            assert year in ["2023"], f"unsupported year for {conference} so far"

            if year == "2023":
                accepted_types = ["Poster", "OralPoster"]

        if select:
            index = input(f"Selecting from {accepted_types}: \n"
                          f"(select by numbers and separate by spaces, e.g., `0` | `0 1` | ...)\n").strip().split()
            accepted_types = [accepted_types[int(i)] for i in index]
        print(f"Outputting abstract of {accepted_types} papers of {conference} {year} ...")

    elif conference in cvf:
        source = "cvf"
        if conference == "CVPR":
            assert year in ["2023", "2022", "2021"], f"unsupported year for {conference} so far"
        elif conference == "ICCV":
            assert year in ["2023", "2021"], f"unsupported year for {conference} so far, " \
                                             f"also note that {conference} is biennial"

        print(f"Outputting abstract of papers of {conference} {year} ...")

    elif conference in other:
        source = None
        # The root web containing ECCV papers surprisingly has the same style as cvf.
        if conference == "ECCV":
            source = "cvf"
            assert year in ["2022", "2020"], f"unsupported year for {conference} so far, " \
                                             f"also note that {conference} is biennial"

        print(f"Outputting abstract of papers of {conference} {year} ...")

    return source, accepted_types


def extract_info(conference: str, paper: dict) -> Dict[str, str]:
    """Extract expected information from requested json data.

    :param conference: Exact abbreviation of conference.
    :param paper: Raw information of each paper in json data.
    :return: Extracted information of each paper.
    """

    def _locate_info(_paper_info: dict | list) -> str | list:
        """Locate accurate information position of different conferences.

        :param _paper_info: Candidate expected information of each paper.
        :return: Accurate expected information of each paper.
        """

        return _paper_info['value'] if isinstance(_paper_info, dict) else _paper_info

    def _restrain_bytes(content: str) -> str:
        """Filter out strange bytes.

        :param content: Content to be restrained.
        :return: Restrained content.
        """

        content = re.sub(r"\n+", " ", content)
        content = re.sub(u"[\\x00-\\x08\\x0b\\x0e-\\x1f\\x7f]", "", content)
        return content

    paper_info = dict()
    paper_info['title'] = _restrain_bytes(_locate_info(paper['content']['title']))
    paper_info['authors'] = _restrain_bytes(", ".join(_locate_info(paper['content']['authors'])))

    try:
        paper_info['keywords'] = _restrain_bytes(", ".join(_locate_info(paper['content']['keywords'])))
    except KeyError:
        paper_info['keywords'] = None

    paper_info['abstract'] = _restrain_bytes(_locate_info(paper['content']['abstract']))

    try:
        if conference == "NeurIPS":
            paper_info['tldr'] = _restrain_bytes(_locate_info(paper['content']['TLDR']))
        elif conference == "ICLR":
            paper_info['tldr'] = _restrain_bytes(_locate_info(paper['content']['TL;DR']))
        elif conference == "ICML":
            paper_info['tldr'] = None
    except KeyError:
        paper_info['tldr'] = None

    return paper_info


def parse_openreview(conference: str, year: str, accepted_types: List[str], toc: bool, pdf: bool, wps: bool) -> None:
    """Parse OpenReview style paper abstracts from the requested json data to a docx file.

    :param conference: Exact abbreviation of conference.
    :param year: Year of conference.
    :param accepted_types: Accepted types of conference.
    :param toc: Indicates that whether table of contents is required.
    :param pdf: Indicates that whether additional pdf is required.
    :param wps: Use WPS to convert pdf.
    """

    def _json_logits(_json_data_path: str, _url: str) -> dict:
        if os.path.exists(_json_data_path):
            with open(_json_data_path, 'r') as _f:
                _json_data = json.load(_f)
        else:
            _json_data = requests.get(_url).json()
            print('json data', _json_data)
            _paper_counts = int(_json_data['count'])
            _default_offset = 1000
            if _paper_counts > _default_offset:
                for _offset in range(_default_offset, _paper_counts, _default_offset):
                    _offset_url = _url + f"&offset={_offset}"
                    _json_data['notes'].extend(requests.get(_offset_url).json()['notes'])

            with open(_json_data_path, 'w') as _f:
                _f.write(json.dumps(_json_data))

        return _json_data

    urls = get_openreview_url(conference, year, accepted_types)
    for accepted_type, url in urls.items():
        document = Document()
        if toc:
            generate_toc(document)
            document.add_page_break()

        json_data = _json_logits(f'logits/{conference}_{year}_{accepted_type}_Abstract.json', url)

        for paper in tqdm(json_data['notes']):
            title = extract_info(conference, paper)['title']
            authors = extract_info(conference, paper)['authors']
            keywords = extract_info(conference, paper)['keywords']
            abstract = extract_info(conference, paper)['abstract']
            tldr = extract_info(conference, paper)['tldr']

            document.add_heading(title, level=2)
            bold_prefix(document, "Authors: ", authors)
            bold_prefix(document, "Keywords: ", keywords)
            bold_prefix(document, "TLDR: ", tldr)
            bold_prefix(document, "Abstract: ", abstract)

        if not os.path.exists('results'):
            os.makedirs('results')
        suffix = "_TOC" if toc else ""
        dst_path = f'results/docx{suffix.lower()}'
        if not os.path.exists(dst_path):
            os.makedirs(dst_path)

        save_path = f"{dst_path}/{int(json_data['count'])}_{accepted_type}_{conference}_{year}_Abstract{suffix}.docx"
        document.save(save_path)
        if toc:
            update_field(os.path.abspath(save_path), wps)
        if pdf:
            generate_pdf(os.path.abspath(save_path), wps)


def parse_cvf(conference: str, year: str, toc: bool, pdf: bool, wps: bool) -> None:
    """Parse CVF style paper abstracts from the requested json data to a docx file.

    :param conference: Exact abbreviation of conference.
    :param year: Year of conference.
    :param toc: Indicates that whether table of contents is required.
    :param pdf: Indicates that whether additional pdf is required.
    :param wps: Use WPS to convert pdf.
    """

    with open(f'logits/{conference}_{year}_Abstract.json', 'rb') as f:
        json_data = json.load(f)

    suffix = "_TOC" if toc else ""
    save_docx_dir = f'results/docx{suffix.lower()}'
    if not os.path.exists(save_docx_dir):
        os.makedirs(save_docx_dir)

    save_docx_path = f"{save_docx_dir}/{len(json_data)}_{conference}_{year}_Abstract{suffix}.docx"
    write_docx = True
    if os.path.exists(save_docx_path):
        overwrite = input("Docx file has existed, overwrite it? (y/n):\n")
        write_docx = False if overwrite == "n" else True

    if write_docx:
        document = Document()
        if toc:
            generate_toc(document)
            document.add_page_break()

        for paper in tqdm(json_data):
            document.add_heading(paper['title'], level=2)
            bold_prefix(document, "Authors: ", paper['authors'])

            # The root web containing ECCV papers surprisingly has the same style as cvf,
            # so seamlessly borrows this function.
            # ECCV paper abstracts have double quotation marks.
            if conference == "ECCV":
                paper['abstract'] = paper['abstract'].strip(r'\"')

            bold_prefix(document, "Abstract: ", paper['abstract'])

        document.save(save_docx_path)
        if toc:
            update_field(os.path.abspath(save_docx_path), wps)

    if pdf:
        generate_pdf(os.path.abspath(save_docx_path), wps)


def bold_prefix(document: Document, prefix: str, content: str) -> None:
    """Emphasize the prefix with the bold style.

    :param document: Document object.
    :param prefix: Prefix to be bold.
    :param content: Main content.
    """

    if content is not None:
        p = document.add_paragraph()
        p.add_run(prefix).bold = True
        p.add_run(content)


def get_openreview_url(conference: str, year: str, accepted_types: List[str]) -> Dict[str, str]:
    """Get OpenReview url given the accepted type.

    :param conference: Exact abbreviation of conference.
    :param year: Year of conference.
    :param accepted_types: Accepted types of conference.
    :return: Accepted_type and OpenReview url.
    """

    url = dict()
    for accepted_type in accepted_types:
        url_v1 = f"https://api.openreview.net/" \
                 f"notes?content.venue={conference}+{year}+{accepted_type}" \
                 f"&details=replyCount&invitation={conference}.cc" \
                 f"%2F{year}%2FConference%2F-%2FBlind_Submission"

        url_v2 = f"https://api2.openreview.net/" \
                 f"notes?content.venue={conference}+{year}+{accepted_type}" \
                 f"&details=replyCount%2Cpresentation&domain={conference}.cc%2F{year}%2FConference"

        if conference == "NeurIPS":
            if year == "2023":
                url[accepted_type] = url_v2
            else:
                url[accepted_type] = url_v1

        elif conference == "ICLR":
            url[accepted_type] = url_v1
            if year == "2024":
                url[accepted_type] = url_v2.replace('+', '%20')
                print(url[accepted_type])
            elif year == "2023":
                # adapt to qualified web format
                web_accepted_type = accepted_type.replace('_', '+').replace('%', '%25')
                url[accepted_type] = url_v1.replace(f"{accepted_type}", web_accepted_type)

        elif conference == "ICML":
            if year == "2023":
                url[accepted_type] = url_v2

    return url


def generate_toc(document: Document) -> None:
    """Generate table of contents.

    :param document: Document object.
    """

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Table Of Contents").bold = True

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_element = run._r  # noqa
    r_element.extend([fldChar, instrText, fldChar2, fldChar3])


def update_field(docx_path: str, wps: bool) -> None:
    """Update field of docx file to display table of contents

    :param docx_path: Path of docx file containing field to be updated.
    :param wps: Use WPS to update table of contents.
    """

    import win32com.client
    print("Updating table of contents ... (wait a little longer)")
    if wps:
        word = win32com.client.DispatchEx("kwps.Application")
    else:
        word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_path)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()


def generate_pdf(docx_path: str, wps: bool) -> None:
    """Additionally generate pdf based on docx file.

    :param docx_path: Path of docx file to be converted to pdf.
    :param wps: Use WPS to convert pdf.
    """

    import win32com.client
    print("Generating pdf ... (wait a little longer)")

    wdFormatPDF = 17
    pdf_path = docx_path.replace("docx", "pdf")
    if not os.path.exists(Path(pdf_path).parent):
        os.makedirs(Path(pdf_path).parent)
    open(pdf_path, "w").close()

    paper_counts = os.path.basename(docx_path).split("_")[0]
    conference = os.path.basename(docx_path).split("_")[1]
    year = os.path.basename(docx_path).split("_")[2]

    if wps:
        word = win32com.client.DispatchEx("kwps.Application")
    elif int(paper_counts) > 3000:
        print(f"{conference} {year} has {paper_counts} papers, "
              f"the size of which may exceed maximal constraint of MS Word, please try WPS.")
        return None
    else:
        word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

